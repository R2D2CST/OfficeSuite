# Python native libraries
import os

# Third party libraries
from docx import Document

# Self build libraries

from Func.AbstractDocument import AbstractDocument


class Word(AbstractDocument):

    def __init__(self, filePath: str) -> None:

        # Class main attribute file path location
        self.filePath = filePath

        # We validate the given path corresponds with an .docx file type
        self.__validateFiletype()

        # We generate our first document attributes
        try:
            self.__readFile()

        # If file does not exists we build a new one
        except FileNotFoundError:
            self.__createNewFile()
            self.__readFile()
        pass

    def __validateFiletype(self) -> bool:
        """
        Method is private not intended to use outside this class type.
        Raises:
            f: Given file path: {self.filePath} Not a valid word file type.

        Returns:
            bool: True if given path is a valid file path and an word filetype, False otherwise.
        """
        if ".docx" not in self.filePath:
            raise f"Given file path: {self.filePath}\n Not a valid word file type."

        else:
            return True

    def __readFile(self) -> None:
        """ """
        self.document = Document(self.filePath)

        self.paragraphsContent = list()

        for paragraph in self.document.paragraphs:
            self.paragraphsContent.append(paragraph.text)

        self.tablesContent = list()

        # Stores the tables content as a 3D table -> row -> column
        self.tablesContent = list()

        for table in self.document.tables:

            # We wipe clean the table list
            tableList = []

            for row in table.rows:
                # We wipe clean the row values
                rowData = []

                for cell in row.cells:

                    # We clear blank spaces
                    rowData.append(cell.text.strip())

                tableList.append(rowData)
            self.tablesContent.append(tableList)

        pass

    def __createNewFile(self) -> None:
        self.document = Document()
        self.document.save(self.filePath)
        pass

    def printDocumentContent(self) -> None:
        """
        Print the entire content of the Word document, including paragraphs and tables.
        """
        print(f"--- Document Content ---")

        # Print all paragraphs
        print("\n** Paragraphs **")
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():  # Ignore empty paragraphs
                print(paragraph.text)

        # Print all tables
        print("\n** Tables **")
        for tableIndex, table in enumerate(self.document.tables, start=1):
            print(f"\nTable {tableIndex}:")
            for row in table.rows:
                rowData = [cell.text.strip() for cell in row.cells]
                print(
                    "\t" + "\t| ".join(rowData)
                )  # Tab-separated values for better formatting

    pass
